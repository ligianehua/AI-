"""产品手册摄入管线：解析（PDF/Word/TXT/MD）-> 章节切块 -> FTS5 索引 -> 检索

检索与知识库同一套 jieba+FTS5 方案；LLM 消化（手册->FAQ候选）见 digest_to_candidates。
"""
import io
import re

from sqlalchemy import text as sql_text
from sqlalchemy.orm import Session

from ..models import LearningCandidate, ManualChunk, ManualDoc
from . import embedding_service
from .kb_search import segment

CHUNK_SIZE = 500      # 目标块大小（字符）
CHUNK_MAX = 700       # 硬上限，超过强制切
HEADING_RE = re.compile(
    r"^\s*(?:#{1,4}\s+.+|第\s*[一二三四五六七八九十百\d]+\s*[章节部分篇].*|[一二三四五六七八九十]+\s*[、.．].+|\d+(?:\.\d+)*\s*[、.．\s].{0,40})\s*$")


class ManualParseError(Exception):
    pass


# ---------- 解析 ----------

def _decode_text(data: bytes) -> str:
    for enc in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


OCR_MAX_PAGES = 20


def _ocr_pdf(data: bytes) -> list[tuple[str, str]]:
    """扫描版 PDF：pymupdf 渲染每页为图片 -> 视觉模型 OCR 提取文字"""
    import base64

    import fitz  # pymupdf

    from ..llm.anthropic_client import get_llm_client

    llm = get_llm_client()
    doc = fitz.open(stream=data, filetype="pdf")
    total_pages = doc.page_count
    pairs: list[tuple[str, str]] = []
    pages = min(total_pages, OCR_MAX_PAGES)
    for page_no in range(pages):
        pix = doc[page_no].get_pixmap(dpi=150)
        img_b64 = base64.b64encode(pix.tobytes("png")).decode()
        try:
            text = llm.ocr_image(img_b64, "image/png")
        except Exception:
            text = ""
        section = f"第{page_no + 1}页"
        for raw_line in (text or "").split("\n"):
            line = raw_line.strip()
            if not line:
                continue
            if HEADING_RE.match(line) and len(line) <= 50:
                section = line
                continue
            pairs.append((section, line))
    doc.close()
    if total_pages > OCR_MAX_PAGES:
        pairs.append(("说明", f"（本手册共 {total_pages} 页，OCR 仅处理前 {OCR_MAX_PAGES} 页）"))
    return pairs


def _parse_pdf(data: bytes) -> list[tuple[str, str]]:
    """返回 [(section, paragraph)]；文字层为空时自动走 OCR（扫描版）"""
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    pairs: list[tuple[str, str]] = []
    section = ""
    for page_no, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if not text.strip():
            continue
        for raw_line in text.split("\n"):
            line = raw_line.strip()
            if not line:
                continue
            if HEADING_RE.match(line) and len(line) <= 50:
                section = line
                continue
            pairs.append((section or f"第{page_no}页", line))
    if not pairs:
        # 扫描版/纯图片 PDF -> OCR 兜底
        pairs = _ocr_pdf(data)
        if not pairs:
            raise ManualParseError("PDF 未提取到文字，OCR 也未识别出内容（请检查图片清晰度或运行模式是否为真实API）")
    return pairs


def _parse_docx(data: bytes) -> list[tuple[str, str]]:
    import docx
    doc = docx.Document(io.BytesIO(data))
    pairs: list[tuple[str, str]] = []
    section = ""
    for p in doc.paragraphs:
        line = p.text.strip()
        if not line:
            continue
        style = (p.style.name or "").lower()
        if "heading" in style or "标题" in style or (HEADING_RE.match(line) and len(line) <= 50):
            section = line
            continue
        pairs.append((section, line))
    for table in doc.tables:  # 表格逐行转文本
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                pairs.append((section, " | ".join(cells)))
    if not pairs:
        raise ManualParseError("Word 文档没有可提取的正文内容")
    return pairs


def _parse_text(data: bytes) -> list[tuple[str, str]]:
    content = _decode_text(data)
    pairs: list[tuple[str, str]] = []
    section = ""
    for raw_line in content.split("\n"):
        line = raw_line.strip()
        if not line:
            continue
        if HEADING_RE.match(line) and len(line) <= 50:
            section = re.sub(r"^#+\s*", "", line)
            continue
        pairs.append((section, line))
    if not pairs:
        raise ManualParseError("文件没有可提取的正文内容")
    return pairs


def parse_file(filename: str, data: bytes) -> tuple[str, list[tuple[str, str]]]:
    """返回 (file_type, [(section, paragraph)])"""
    ext = (filename.rsplit(".", 1)[-1] if "." in filename else "").lower()
    if ext == "pdf":
        return "pdf", _parse_pdf(data)
    if ext in ("docx", "doc"):
        if ext == "doc":
            raise ManualParseError("暂不支持旧版 .doc，请另存为 .docx 后上传")
        return "docx", _parse_docx(data)
    if ext in ("txt", "md", "markdown", ""):
        return ("md" if ext.startswith("m") else "txt"), _parse_text(data)
    raise ManualParseError(f"不支持的文件类型 .{ext}（支持 PDF / DOCX / TXT / MD）")


# ---------- 切块 ----------

def build_chunks(pairs: list[tuple[str, str]]) -> list[tuple[str, str]]:
    """把 (章节, 段落) 序列打包成 ~CHUNK_SIZE 字的块，章节变化即断块"""
    chunks: list[tuple[str, str]] = []
    cur_section, buf = None, ""

    def flush():
        nonlocal buf
        if buf.strip():
            chunks.append((cur_section or "", buf.strip()))
        buf = ""

    for section, para in pairs:
        if section != cur_section:
            flush()
            cur_section = section
        candidate = (buf + "\n" + para).strip() if buf else para
        if len(candidate) >= CHUNK_SIZE:
            if len(candidate) <= CHUNK_MAX:
                buf = candidate
                flush()
            else:
                flush()
                # 超长段按句切
                buf = ""
                for sent in re.split(r"(?<=[。！？!?；;])", para):
                    if len(buf) + len(sent) > CHUNK_SIZE and buf:
                        flush()
                    buf += sent
        else:
            buf = candidate
    flush()
    return chunks


# ---------- 入库与索引 ----------

def ingest(db: Session, filename: str, data: bytes) -> ManualDoc:
    file_type, pairs = parse_file(filename, data)
    chunks = build_chunks(pairs)
    if not chunks:
        raise ManualParseError("切块结果为空")
    title = re.sub(r"\.[^.]+$", "", filename)
    doc = ManualDoc(filename=filename, title=title, file_type=file_type,
                    chunk_count=len(chunks),
                    char_count=sum(len(c) for _, c in chunks))
    db.add(doc)
    db.flush()
    # 批量向量化（每批64块；embeddings 不可用时静默跳过）
    vecs = embedding_service.embed_texts(
        [f"{s}\n{c}" for s, c in chunks]) if len(chunks) <= 64 else None
    if vecs is None and len(chunks) > 64:
        vecs = []
        for i in range(0, len(chunks), 64):
            batch = embedding_service.embed_texts(
                [f"{s}\n{c}" for s, c in chunks[i:i + 64]])
            if batch is None:
                vecs = None
                break
            vecs.extend(batch)
    for seq, (section, content) in enumerate(chunks, start=1):
        chunk = ManualChunk(doc_id=doc.id, seq=seq, section=section[:200], content=content)
        if vecs:
            chunk.embedding = embedding_service.to_json(vecs[seq - 1])
        db.add(chunk)
        db.flush()
        db.execute(sql_text(
            "INSERT INTO manual_fts (chunk_id, section_seg, content_seg) "
            "VALUES (:cid, :s, :c)"),
            {"cid": chunk.id, "s": segment(section), "c": segment(content)})
    db.commit()
    return doc


def delete_doc(db: Session, doc: ManualDoc):
    chunk_ids = [c.id for c in db.query(ManualChunk).filter(ManualChunk.doc_id == doc.id)]
    if chunk_ids:
        db.execute(sql_text(
            f"DELETE FROM manual_fts WHERE chunk_id IN ({','.join(map(str, chunk_ids))})"))
    db.query(ManualChunk).filter(ManualChunk.doc_id == doc.id).delete()
    db.delete(doc)
    db.commit()


# ---------- 检索 ----------

def search(db: Session, query: str, top_k: int = 3) -> list[dict]:
    """混合检索手册块（FTS + 向量，RRF融合），返回 [{doc, section, content, score}]"""
    fts_ids, fts_scores = [], {}
    seg = segment(query)
    if seg:
        tokens = seg.split()
        strong = [t for t in tokens if len(t) >= 2]
        match_expr = " OR ".join(f'"{t}"' for t in (strong or tokens))
        try:
            rows = db.execute(sql_text(
                "SELECT chunk_id, bm25(manual_fts, 3.0, 1.0) AS score "
                "FROM manual_fts WHERE manual_fts MATCH :m ORDER BY score LIMIT :k"),
                {"m": match_expr, "k": top_k * 4}).fetchall()
            fts_ids = [r[0] for r in rows]
            fts_scores = {r[0]: -float(r[1]) for r in rows}
        except Exception:
            pass

    vec_ids, vec_sims = [], {}
    qvec = embedding_service.embed_one(query)
    if qvec:
        cands = db.query(ManualChunk.id, ManualChunk.embedding) \
            .filter(ManualChunk.embedding.isnot(None)).all()
        ranked = embedding_service.vector_rank(qvec, cands, top_k=8)
        vec_ids = [cid for cid, _ in ranked]
        vec_sims = dict(ranked)

    merged = embedding_service.rrf_merge([fts_ids, vec_ids], top_k=top_k) \
        if vec_ids else fts_ids[:top_k]
    out = []
    for chunk_id in merged:
        chunk = db.query(ManualChunk).get(chunk_id)
        if chunk is None:
            continue
        score = fts_scores.get(chunk_id) or round(vec_sims.get(chunk_id, 0) * 10, 3)
        out.append({"doc": chunk.doc.title, "section": chunk.section,
                    "content": chunk.content, "score": round(score, 3)})
    return out


def backfill_embeddings(db: Session) -> int:
    """给缺失向量的手册块补齐 embedding（重建索引时调用）"""
    missing = db.query(ManualChunk).filter(ManualChunk.embedding.is_(None)).all()
    done = 0
    for i in range(0, len(missing), 64):
        batch = missing[i:i + 64]
        vecs = embedding_service.embed_texts(
            [f"{c.section}\n{c.content}" for c in batch])
        if vecs is None:
            break
        for c, v in zip(batch, vecs):
            c.embedding = embedding_service.to_json(v)
            done += 1
    db.commit()
    return done


# ---------- LLM 消化：手册 -> FAQ 候选 ----------

def digest_to_candidates(doc_id: int):
    """后台任务：让 LLM 通读手册块，提炼 FAQ 候选进学习审核队列"""
    from ..database import SessionLocal
    from ..llm.anthropic_client import get_llm_client

    db = SessionLocal()
    try:
        doc = db.query(ManualDoc).get(doc_id)
        if not doc:
            return
        chunks = db.query(ManualChunk).filter(ManualChunk.doc_id == doc.id) \
            .order_by(ManualChunk.seq).all()
        # 控制总量：取前 ~12000 字
        parts, total = [], 0
        for c in chunks:
            piece = (f"【{c.section}】\n" if c.section else "") + c.content
            if total + len(piece) > 12000:
                break
            parts.append(piece)
            total += len(piece)
        manual_text = "\n\n".join(parts)

        llm = get_llm_client()
        candidates = llm.digest_manual(doc.title, manual_text)

        created = 0
        for item in candidates[:12]:
            q = str(item.get("question", "")).strip()
            a = str(item.get("suggested_answer", "")).strip()
            if len(q) < 4 or len(a) < 10:
                continue
            db.add(LearningCandidate(
                type="faq", question=q, suggested_answer=a,
                category=str(item.get("category", "产品使用")),
                confidence=float(item.get("confidence", 0.8)),
                review_note=f"来自手册《{doc.title}》"))
            created += 1
        db.commit()
        return created
    finally:
        db.close()
