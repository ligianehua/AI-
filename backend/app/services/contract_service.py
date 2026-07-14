"""M10 合同处理：上传 → 异步抽取+审查；标准合同草稿 docx 生成。RBAC 在本层强制。

红线：AI 输出仅为提示，不构成法律意见（生成文档与 UI 均有声明）。
"""

import io
import uuid
from datetime import UTC, datetime
from functools import lru_cache
from pathlib import Path
from typing import Any

import yaml
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.exceptions import DomainError, NotFoundError
from app.models.account import Account
from app.models.contract import Contract
from app.models.enums import ContractStatus
from app.models.user import User
from app.services.base import BaseService
from app.services.knowledge_service import UPLOAD_DIR as KNOWLEDGE_UPLOAD_DIR
from app.services.opportunity_service import opportunity_service
from app.tasks import dispatcher

CONTRACT_UPLOAD_DIR = KNOWLEDGE_UPLOAD_DIR / "contracts"
ALLOWED_SUFFIXES = {".txt", ".md", ".docx"}
MAX_TEXT_CHARS = 30_000  # 送入 LLM 的合同文本上限（超长截断，防 token 爆炸）

_RISK_RULES_FILE = Path(__file__).parent / "contract_risk_rules.yaml"

DRAFT_DISCLAIMER = "本合同由 AI 销售助手生成的草稿，仅供参考，正式签署前须经法务审核。"


@lru_cache
def get_risk_checklist() -> list[dict[str, str]]:
    data = yaml.safe_load(_RISK_RULES_FILE.read_text(encoding="utf-8"))
    return list(data["checklist"])


class ContractService(BaseService[Contract]):
    model = Contract
    sortable_fields = frozenset({"created_at", "name", "status"})


contract_service = ContractService()


async def upload_contract(
    session: AsyncSession, actor: User, filename: str, content: bytes
) -> Contract:
    suffix = Path(filename).suffix.lower()
    if suffix not in ALLOWED_SUFFIXES:
        raise DomainError(f"不支持的文件类型 {suffix}（支持 txt/md/docx）")
    if len(content) > 10 * 1024 * 1024:
        raise DomainError("文件不能超过 10MB")

    contract = Contract(
        name=Path(filename).stem[:200],
        status=ContractStatus.PROCESSING,
        owner_id=actor.id,
    )
    session.add(contract)
    await session.commit()
    await session.refresh(contract)

    CONTRACT_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    (CONTRACT_UPLOAD_DIR / f"{contract.id}{suffix}").write_bytes(content)

    enqueued = await dispatcher.enqueue("process_contract_task", str(contract.id))
    if not enqueued:
        contract.status = ContractStatus.FAILED
        contract.error_msg = "任务投递失败，请点重试"
        await session.commit()
    return contract


def contract_file_path(contract_id: uuid.UUID) -> Path | None:
    for suffix in ALLOWED_SUFFIXES:
        candidate = CONTRACT_UPLOAD_DIR / f"{contract_id}{suffix}"
        if candidate.exists():
            return candidate
    return None


async def reprocess_contract(session: AsyncSession, actor: User, contract_id: uuid.UUID) -> None:
    contract = await contract_service.get(session, actor, contract_id)
    if contract_file_path(contract.id) is None:
        raise DomainError("合同源文件缺失，请删除后重新上传")
    contract.status = ContractStatus.PROCESSING
    contract.error_msg = None
    await session.commit()
    await dispatcher.enqueue("process_contract_task", str(contract.id))


async def delete_contract(session: AsyncSession, actor: User, contract_id: uuid.UUID) -> None:
    contract = await contract_service.get(session, actor, contract_id)
    contract.deleted_at = datetime.now(UTC)
    await session.commit()


# ---------- 标准合同草稿生成（python-docx 代码生成，不依赖模板文件） ----------


async def generate_draft_docx(
    session: AsyncSession,
    actor: User,
    opportunity_id: uuid.UUID,
    payment_terms: str | None = None,
) -> tuple[str, bytes]:
    """从商机生成标准销售合同草稿。返回 (文件名, docx bytes)。"""
    opp = await opportunity_service.get(session, actor, opportunity_id)
    account = await session.scalar(select(Account).where(Account.id == opp.account_id))
    if account is None:
        raise NotFoundError("商机所属客户不存在")

    from docx import Document

    doc: Any = Document()
    doc.add_paragraph(DRAFT_DISCLAIMER)
    doc.add_heading("销售服务合同（草稿）", level=0)

    today = datetime.now(UTC).date().isoformat()
    doc.add_paragraph(f"甲方（客户）：{account.name}")
    doc.add_paragraph("乙方（服务方）：____________________")
    doc.add_paragraph(f"签署日期：{today}（草稿生成日，以实际签署为准）")

    doc.add_heading("一、服务内容", level=1)
    doc.add_paragraph(f"乙方向甲方提供「{opp.name}」相关的产品与服务，具体范围以附件报价单为准。")

    doc.add_heading("二、合同金额", level=1)
    doc.add_paragraph(f"合同总金额为人民币 {opp.amount:,.2f} 元（含税）。")

    doc.add_heading("三、付款方式", level=1)
    doc.add_paragraph(
        payment_terms
        or "合同签署后 7 个工作日内支付合同总额的 50%；验收合格后 7 个工作日内支付剩余 50%。"
    )

    doc.add_heading("四、服务期限", level=1)
    expected = opp.expected_close_date.isoformat() if opp.expected_close_date else "____"
    doc.add_paragraph(f"自合同签署之日起 12 个月（预计启动日：{expected}）。")

    doc.add_heading("五、验收标准", level=1)
    doc.add_paragraph(
        "双方按附件约定的功能清单验收；甲方应在乙方交付后 10 个工作日内完成验收，逾期视为验收通过。"
    )

    doc.add_heading("六、违约责任", level=1)
    doc.add_paragraph("任一方违约，守约方有权要求违约方支付合同总额 10% 的违约金；双方责任对等。")

    doc.add_heading("七、保密条款", level=1)
    doc.add_paragraph("双方对合作中知悉的对方商业信息负有保密义务，保密期自合同终止后 2 年。")

    doc.add_heading("八、争议解决", level=1)
    doc.add_paragraph(
        "因本合同产生的争议，双方友好协商解决；协商不成的，提交乙方所在地人民法院管辖。"
    )

    doc.add_heading("九、签署", level=1)
    doc.add_paragraph("甲方（盖章）：____________　　乙方（盖章）：____________")
    doc.add_paragraph("日期：____________　　　　　　日期：____________")
    doc.add_paragraph("")
    doc.add_paragraph(DRAFT_DISCLAIMER)

    buf = io.BytesIO()
    doc.save(buf)
    return f"{account.name}-{opp.name}-合同草稿.docx", buf.getvalue()
